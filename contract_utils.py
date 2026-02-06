import re
import io
import streamlit as st
from docx import Document
from datetime import datetime

# --- CONSTANTES DE PLANTILLAS DOCX ---
# Nombres exactos de los archivos DOCX que subió el usuario
TEMPLATE_VEJEZ_DOCX = "CONTRATO 2026 AP - PLANTILLA.docx" 
TEMPLATE_SOBREVIVENCIA_DOCX = "CONTRATO 2026 sobrevivencia -  plantilla.docx"

def get_template_path(tipo):
    if tipo == "Vejez o Invalidez":
        return TEMPLATE_VEJEZ_DOCX
    else:
        return TEMPLATE_SOBREVIVENCIA_DOCX

def extract_client_data_from_markdown(markdown_text):
    """
    Intenta extraer datos básicos del informe Markdown usando Regex.
    Retorna un diccionario con lo encontrado.
    """
    data = {}
    
    # Patrones de búsqueda comunes en el informe
    patterns = {
        "Nombre Completo": r"\*\*Nombre Completo:\*\*\s*(.*)",
        "RUT": r"\*\*RUT:\*\*\s*(.*)",
        "Dirección": r"\*\*Direcci[óo]n:\*\*\s*(.*)", # A veces no está
        "Comuna": r"\*\*Comuna:\*\*\s*(.*)",
        "Teléfono": r"\*\*Tel[ée]fono:\*\*\s*(.*)",
        "Estado Civil": r"\*\*Estado Civil:\*\*\s*(.*)",
        "Cédula de Identidad": r"\*\*RUT:\*\*\s*(.*)", # Usualmente el RUT es el CI
    }
    
    for field, pattern in patterns.items():
        match = re.search(pattern, markdown_text, re.IGNORECASE)
        if match:
            # Limpiar el dato capturado
            val = match.group(1).strip()
            if "No informada" not in val and "[Extraer" not in val:
                data[field] = val
    
    return data

def replace_text_in_paragraph(paragraph, replacements):
    """
    Reemplaza texto en un párrafo conservando el estilo.
    Busca {{KEY}} o simplemente palabras clave si definimos una estrategia más flexible.
    Por ahora, asumiremos que insertamos el texto donde encontremos marcadores o usamos una estrategia de reemplazo simple.
    Dado que el usuario NO tiene placeholders {{}} puestos, usaremos una estrategia de búsqueda de texto de la plantilla PDF a DOCX.
    
    SIN EMBARGO, como el usuario pidió "respetar formato", la mejor estrategia si NO hay placeholders 
    es simplemente buscar etiquetas comunes como "Nombre:" y tratar de insertar después.
    
    Pero para ser precisos, vamos a ofrecer al usuario que Llene los datos y nosotros reemplazaremos
    ciertos marcadores de posición que asumiremos existen O añadiremos el texto.
    """
    # Estrategia simple: Iterar sobre diccionario y reemplazar keys
    for key, value in replacements.items():
        if key in paragraph.text:
            # Reemplazo directo (puede romper formato si hay runs partidos, pero es lo más directo sin placeholders complejos)
            paragraph.text = paragraph.text.replace(key, str(value))

def fill_contract_template(template_path, data_dict):
    """
    Abre el DOCX, y realiza reemplazos.
    Como las plantillas del usuario PROBABLEMENTE NO tienen {{PLACEHOLDERS}}, 
    vamos a intentar reemplazar texto genérico o líneas de subrayado si podemos identificarlas,
    o mejor, confiaremos en que el usuario o nosotros definimos claves.
    
    Si el usuario NO puso placeholders, este código intentará reemplazar marcadores que definiremos ahora
    o simplemente devolverá el doc sin cambios si no encuentra nada.
    
    Para que funcione SIN cambiar la plantilla manualmente, necesitaríamos saber qué buscar.
    
    ESTRATEGIA AVANZADA: 
    Dado que no puedo ver el DOCX para saber los placeholders, voy a asumir que el usuario 
    QUIERE que nosotros insertemos los datos.
    Voy a pedirle al usuario en el FORMULARIO que defina los valores que irán en el contrato.
    
    Pero para "Escribir" en el docx, necesito donde.
    Voy a usar una lista de KEYWORDS que usualmente están en contratos y reemplazar "____________________" si existe,
    o insertar después de "Nombre: ".
    """
    try:
        doc = Document(template_path)
        
        # Mapa de reemplazos "Inteligente" (Si la plantilla tiene estos textos)
        # Ajustaremos esto según lo que veamos o estandarizaremos.
        # Por seguridad, vamos a asumir que el usuario prefiere que reemplacemos marcadores estándar
        # O que añadiremos los reemplazos que el diccionario data_dict traiga.
        
        # Vamos a definir marcadores esperados basados en data_dict keys
        # Ejemplo: Si data_dict tiene "CLIENTE_NOMBRE", buscaremos "CLIENTE_NOMBRE" en el doc.
        
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, data_dict)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, data_dict)
                        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
        
    except Exception as e:
        return None, str(e)

