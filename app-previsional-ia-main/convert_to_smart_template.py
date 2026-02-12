from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def convert_template(input_path, output_path, is_survivorship=False):
    print(f"Converting {input_path} -> {output_path}...")
    doc = Document(input_path)
    
    # 1. Beneficiary Table Logic (Only if survivorship/placeholders found)
    insert_point_index = -1
    paragraphs_to_remove = []
    
    for i, p in enumerate(doc.paragraphs):
        text = p.text
        if "{NOMBRE BENEFICIARIO}" in text:
            if insert_point_index == -1:
                insert_point_index = i
            paragraphs_to_remove.append(p)

    if insert_point_index != -1:
        print(f"  Found beneficiary placeholders at {insert_point_index}. Replacing with table.")
        # Remove old
        for p in paragraphs_to_remove:
             p._element.getparent().remove(p._element)
             
        # Insert Table
        if insert_point_index < len(doc.paragraphs):
            ref_p = doc.paragraphs[insert_point_index]
            ref_element = ref_p._element
        else:
            ref_element = None
            
        table = doc.add_table(rows=2, cols=3)
        # remove style assignment
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Nombre Completo"
        hdr_cells[1].text = "RUT"
        hdr_cells[2].text = "Parentesco"
        
        row = table.rows[1]
    
        # We use standard for loop spanning cells to repeat row
        row.cells[0].text = "{% for b in beneficiaries %}{{ b.nombre }}"
        row.cells[1].text = "{{ b.rut }}"
        row.cells[2].text = "{{ b.parentesco }}{% endfor %}"

        tbl, p_elem = table._element, ref_element
        if p_elem is not None:
            p_elem.addprevious(tbl)
        else:
            doc._body._element.append(tbl)
    else:
        print("  No beneficiary placeholders found. Skipping table insertion.")

    # 2. Global Replacements
    replacements_map = {
        # Common
        "{NOMBRE AFILIADO}": "{{ nombre_afiliado }}",
        "{RUT AFILIADO}": "{{ rut_afiliado }}",
        "{DIRECCIÓN}": "{{ direccion_afiliado }}",
        "{COMUNA}": "{{ comuna_afiliado }}",
        "{CIUDAD}": "{{ ciudad_afiliado }}",
        "{TELEFONO}": "{{ telefono_afiliado }}",
        "{CELULAR}": "{{ celular_afiliado }}",
        "{CORREO ELECTRÓNICO}": "{{ correo_afiliado }}",
        "{ESTADO CIVIL AFILIADO}": "{{ estado_civil_afiliado }}",
        "{FECHA DE NACIMIENTO AFILIADO}": "{{ fecha_nacimiento_afiliado }}",
        "{OFICIO AFILIADO}": "{{ oficio_afiliado }}",
        "{AFP DE ORIGEN}": "{{ afp_origen }}",
        "{SISTEMA DE SALUD}": "{{ sistema_salud }}",
        "{TIPO DE PENSIÓN}": "{{ tipo_pension }}",
        "{FECHA}": "{{ fecha_actual }}",
        
        # Survivorship
        "{NOMBRE CAUSANTE}": "{{ nombre_causante }}",
        "{RUT CAUSANTE}": "{{ rut_causante }}",
        "{NOMBRE CONSULTANTE}": "{{ nombre_consultante }}",
        "{RUT CONSULTANTE}": "{{ rut_consultante }}",
        "{PROFESIÓN CONSULTANTE}": "{{ profesion_consultante }}",
        "{ESTADO CIVIL CONSULTANTE}": "{{ estado_civil_consultante }}",
        "{FECHA DE NACIMIENTO CONSULTANTE}": "{{ fecha_nacimiento_consultante }}",

        # Fallbacks / Cleanup
        "{NOMBRE}": "{{ nombre_afiliado }}",
        "{RUT}": "{{ rut_afiliado }}",
        "{PARENTESCO}": "", # Should be handled by table, but clean up if stray
        "{NOMBRE BENEFICIARIO}": "", 
        "{RUT BENEFICIARIO}": "",
        "{FECHA NACIMIENTO  BENEFICIARIO}": "", # Double space
    }
    
    for p in doc.paragraphs:
        for old, new in replacements_map.items():
            if old in p.text:
                p.text = p.text.replace(old, new)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old, new in replacements_map.items():
                         if old in p.text:
                            p.text = p.text.replace(old, new)

    doc.save(output_path)
    print(f"  Saved {output_path}")

if __name__ == "__main__":
    convert_template("CONTRATO_FIXED_SURVIVORSHIP.docx", "CONTRATO_SMART_SURVIVORSHIP.docx", is_survivorship=True)
    convert_template("CONTRATO_FIXED_OLD_AGE.docx", "CONTRATO_SMART_OLD_AGE.docx", is_survivorship=False)
