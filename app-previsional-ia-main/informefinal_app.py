import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io
import google.generativeai as genai
from datetime import datetime
import re

# --- IMPORTS PARA OCR ---
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from contract_utils import get_contract_template_path, extract_contract_data, generate_contract_docx




# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(layout="wide", page_title="Asesor Previsional IA")

st.sidebar.info("🤖 Asistente de Asesoría Previsional IA")
st.sidebar.divider()
st.sidebar.subheader("Modificar Informe")

# Caja de texto para modificar el informe
instrucciones_mod = st.sidebar.text_area(
    "Indicaciones de Modificación",
    help="Usa esta caja para pedirle a la IA que refine el informe (ej. 'Elimina los puntos 1 y 2 de la nota', 'Acorta la sección 6', 'Cambia el tono a más formal').",
    key="instrucciones_mod"
)

# -------------------------------


# --- 2. FUNCIONES DE LECTURA Y IA ---

@st.cache_data
def leer_pdfs_cargados(files):
    """
    Lee el texto de múltiples archivos PDF.
    Si una página parece escaneada, aplica OCR automáticamente.
    """
    contexto_completo = ""
    st.write("Archivos recibidos para análisis:")
    
    for file in files:
        st.caption(f"- {file.name}")
        try:
            full_text = ""
            # Abrir el PDF en memoria con PyMuPDF (fitz)
            doc = fitz.open(stream=io.BytesIO(file.read()), filetype="pdf")
            
            for i, page in enumerate(doc):
                page_num = i + 1
                
                # 1. Intentar extracción de texto digital
                # sort=True intenta ordenar el texto por posición (útil para tablas)
                text = page.get_text("text", sort=True)
                
                # 2. Heurística: Si el texto es muy corto, probablemente es escaneado
                if len(text.strip()) < 150: # Umbral de 150 caracteres
                    st.warning(f"Página {page_num} de {file.name} parece escaneada. Iniciando OCR... (esto puede tardar)")
                    
                    # 3. Renderizar la página como imagen (300 DPI)
                    zoom = 300 / 72  # 300 DPI / 72 DPI (default)
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # 4. Convertir a formato PIL (Pillow)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # 5. Usar Tesseract para OCR en español
                    try:
                        # 'spa' = Spanish
                        ocr_text = pytesseract.image_to_string(img, lang='spa')
                        full_text += f"\n\n--- PÁGINA {page_num} ({file.name}) [Texto extraído por OCR] ---\n\n{ocr_text}"
                    except Exception as ocr_error:
                        st.error(f"Error de OCR en página {page_num}. Asegúrate de que Tesseract esté instalado y 'spa' (español) esté disponible. Error: {ocr_error}")
                        full_text += f"\n\n--- PÁGINA {page_num} ({file.name}) [ERROR DE OCR] ---\n\n"
                
                else:
                    # Es un PDF digital, usar el texto extraído
                    full_text += f"\n\n--- PÁGINA {page_num} ({file.name}) [Texto digital] ---\n\n{text}"
            
            contexto_completo += f"\n\n=== INICIO DOCUMENTO: {file.name} ===\n{full_text}\n=== FIN DOCUMENTO: {file.name} ===\n\n"
            doc.close()
        
        except Exception as e:
            st.error(f"Error al leer {file.name}: {e}")
    return contexto_completo

# === PROMPT PASO 1: ANÁLISIS (SECCIONES 1-5) - SIN COMISIONES AFP ---
PROMPT_ANALISIS = """
Eres un Asesor Previsional experto y senior, con profundo conocimiento del sistema de pensiones chileno (AFP, SCOMP, PGU, APV, etc.).
Tu tarea es analizar TODOS los documentos de antecedentes que te entregaré (SCOMP, Certificado de Saldo, etc.) y generar un **Informe de Análisis** que contenga ÚNICAMENTE las secciones 1 a 5.
REGLAS IMPORTANTES:
1.  **Actúa como un experto:** Tu tono debe ser profesional y claro.
2.  **Cíñete a los datos:** No inventes información. Si un dato no se encuentra en los documentos (ej. Fecha de Nacimiento), debes indicarlo explícitamente (ej: "Fecha de Nacimiento: No informada en los documentos").
3.  **Calcula cuando se pida:** Para las Rentas Vitalicias Aumentadas, DEBES calcular los montos aumentados (Pensión Aumentada UF/$, Pensión Líquida Aumentada) basándote en la "pensión base" que encuentres en el SCOMP.
4.  **Usa Markdown:** Estructura tu respuesta usando Markdown (títulos, negritas, tablas).
5.  **Fecha del Informe:** {FECHA_HOY}
6.  **NO INCLUYAS la Sección 6 (Recomendación Final).** Termina el informe después de la Sección 5.
7.  **Formato de Títulos:** Usa '##' para Secciones (ej. ## 1) Antecedentes) y '###' para Subsecciones (ej. ### Certificado de Saldos). Usa '####' para los títulos de las modalidades (ej. #### a) Retiro programado).
8.  **NO INCLUIR COMISIONES AFP (SOLO EN TABLA):** En la tabla de Retiro Programado, NO incluyas la columna "Comisión AFP". Sin embargo, SÍ debes considerar la comisión para el cálculo en la "Nota" explicativa debajo de la tabla.
9.  **IMPORTANTE - ALINEACIÓN DE TABLAS:** Al extraer datos de tablas (especialmente SCOMP), ten mucho cuidado de **asociar correctamente cada AFP con SU monto**.
    *   Si el texto extraído muestra primero una lista de AFPs y luego una lista de montos, **compagínalos en el orden en que aparecen**.
    *   Verifica fila por fila. No mezcles la pensión de una AFP con el nombre de otra.
    *   Si hay montos en UF y Pesos, asegúrate de poner cada uno en su columna correcta.
10.  **IMPOTANTE - CHAIN OF THOUGHT (LISTAR MODALIDADES):** Antes de generar el informe, analiza internamente todas las modalidades de pensión presentes en el SCOMP (ej. Renta Vitalicia Inmediata con Retiro, Sin Retiro, Garantizada 120, 240, etc.). Asegúrate de no omitir NINGUNA en el informe final, especialmente las Garantizadas.
---
TEXTO EXTRAÍDO DE LOS DOCUMENTOS DEL CLIENTE (SCOMP, CARTOLAS, ETC.):
{CONTEXTO_DOCUMENTOS}
---
Basado ÚNICAMENTE en los documentos, genera el informe con la siguiente estructura exacta (Secciones 1 a 5):

## Informe final de Asesoría Previsional
### 1) Antecedentes del afiliado y Solicitud de Ofertas
[INSTRUCCIÓN CRÍTICA: Busca específicamente en el documento "Solicitud de Ofertas" para extraer los siguientes datos con mayor precisión. Si no están ahí, búscalos en el SCOMP.]
* **Nombre Completo:** [Extraer]
* **RUT:** [Extraer]
* **Fecha de Nacimiento:** [Extraer]
* **Edad Cumplida (a la fecha actual):** [Calcular o extraer si está]
* **Sexo:** [Extraer]
* **Estado Civil:** [Extraer]
* **Dirección:** [Extraer dirección completa incluyendo Comuna y Ciudad desde Solicitud de Ofertas]
* **Correo Electrónico:** [Extraer e-mail desde Solicitud de Ofertas]
* **Teléfono/Celular:** [Extraer]
* **AFP de Origen:** [Extraer desde Solicitud de Ofertas]
* **Institución de Salud:** [Extraer o poner "No informada"]
* **Fecha Solicitud de Pensión:** [Extraer]
* **Fecha Solicitud de Ofertas:** [Extraer fecha del encabezado del formulario Solicitud de Ofertas]
* **Tipo de Pensión Solicitada:** [Extraer desde Solicitud de Ofertas, ej: Vejez Edad]


#### Datos de Sobrevivencia (Solo si aplica)
[Si el trámite es de Sobrevivencia, extrae lo siguiente del SCOMP / Solicitud de Ofertas]
* **Causante Nombre:** [Nombre del fallecido]
* **Causante RUT:** [RUT del fallecido]
* **Consultante Nombre:** [Nombre de quien solicita]
* **Consultante RUT:** [RUT de quien solicita]

#### Certificado de Saldos
**Descripción:** El saldo total destinado a pensión (Cotizaciones Obligatorias, Fondo [Extraer Fondo]) es de **UF [Extraer Saldo UF]**. Este monto equivale a **$[Extraer Saldo $]**. El valor de la UF utilizado es de **$[Extraer Valor UF]** al **[Extraer Fecha UF]**. Este Certificado se encuentra vigente hasta el día **[Extraer Vigencia Saldo]**.
### 2) Antecedentes del beneficiario
El afiliado declara a la siguiente beneficiaria legal de pensión:
[Generar una tabla Markdown con TODOS los beneficiarios encontrados en el SCOMP. Si no hay, indicar "Sin beneficiarios declarados".]
| Nombre Completo | RUT | Parentesco | Sexo | Invalidez | Fecha de Nacimiento |
| :--- | :--- | :--- | :--- | :--- | :--- |
| [Nombre] | [RUT] | [Parentesco] | [F/M] | [S/N] | [Fecha] |
| [Nombre 2] | [RUT 2] | [Parentesco 2] | [F/M] | [S/N] | [Fecha] |
### 3) Situación previsional
* **Tipo de Pensión Solicitada:** [Extraer, ej: Vejez Edad, Cambio de Modalidad]
* **Saldo para Pensión:** **UF [Extraer Saldo UF]**
* **Modalidades Solicitadas al SCOMP:** [Extraer las modalities que se pidieron, ej: RVIS, RVA 100% 36m]
### 4) Gestiones realizadas
[Describir las gestiones en formato lista o tabla, extrayendo fechas y acciones. Ej:
* **Solicitud de Pensión de Vejez Edad:** Presentada el [Fecha] a AFP [Nombre].
* **Retiro Certificado de Saldos:** Se retira el día [Fecha].
* **Solicitud de Ofertas (SCOMP):** Ingresada el [Fecha], por el Asesor Previsional [Nombre Asesor].]
* **Modalidades Solicitadas:** [Extraer TODAS las modalidades marcadas con 'X' en la Solicitud de Ofertas, incluyendo meses garantizados y cláusulas. Ej: "Retiro Programado", "Renta Vitalicia Inmediata con Condiciones Especiales de Cobertura: 240 meses garantizados"]
### 5) Resultados Scomp
#### a) Retiro programado
**Descripción:** Es una modalidad de pensión que se paga con cargo a la Cuenta de Capitalización Individual del afiliado. La pensión se recalcula anualmente, considerando el saldo remanente, la expectativa de vida del afiliado y de sus beneficiarios, y la rentabilidad del fondo. Por lo tanto, la pensión puede subir o bajar cada año.
**Cuadro de resultados:**
[Generar tabla Markdown con TODAS las AFP del SCOMP]
| AFP | Pensión en UF | Pensión Bruta en $| Descuento 7% Salud$ | Pensión Líquida en $ |
| :--- | :--- | :--- | :--- | :--- |
| [AFP 1] | [uf] | [bruta] | [salud] | [liquida] |
| [AFP 2] | [uf] | [bruta] | [salud] | [liquida] |
| ... | ... | ... | ... | ... |
| [AFP 2] | [uf] | [bruta] | [salud] | [liquida] |
| ... | ... | ... | ... | ... |
| [AFP 2] | [uf] | [bruta] | [salud] | [liquida] |
| ... | ... | ... | ... | ... |
**Nota:** La oferta de Retiro Programado de su AFP de Origen ([Nombre AFP]) es de **[UF] UF** al mes, lo que equivale a una Pensión Bruta de **$[MontoBruto]**. Con el descuento de salud 7% ($[MontoSalud]) y la comisión de administración de la AFP del [Comision]% ($[MontoComision]), la pensión líquida aproximada es de **$[MontoLiquido]** para el primer año.
*(Instrucción: Busca el % de comisión de la AFP de origen en el certificado de saldo o oferta interna. Calcula el monto en pesos [Bruta * %]. Resta Salud y Comisión a la Bruta para obtener la Líquida).*

[INSTRUCCIÓN CLAVE: Si en el SCOMP aparece "Pensión de Referencia Garantizada por ley" (común en Invalidez), AGREGA AQUÍ LA SIGUIENTE SECCIÓN b). Si no, salta a Renta Vitalicia.]

#### b) Pensión de Referencia Garantizada por ley
**Descripción:** Por ley las Compañías de Seguros de Vida indicadas más abajo, garantizan una Pensión de referencia con su saldo obligatorio hasta la fecha de vigencia indicada. El monto de la Pensión garantizada, en renta vitalicia inmediata simple, será el siguiente:
**Cuadro de resultados:**
[Generar tabla con los datos de Referencia Garantizada del SCOMP]

#### [Si hubo sección b, esta es c), sino b)] Renta Vitalicia

**Renta Vitalicia Inmediata Simple**
**Descripción:** Es un contrato con una Compañía de Seguros, donde el afiliado traspasa la totalidad de su saldo para recibir una pensión mensual en UF fija y de por vida. El monto no varía, independiente de la rentabilidad del mercado o de la expectativa de vida.
**Cuadro de resultados (4 mejores ofertas):**
| Compañía de Seguros | Pensión en UF | Pensión Bruta $| Descuento 7% Salud$ | Pensión Líquida $ |
| :--- | :--- | :--- | :--- | :--- |
| [Cia 1] | [uf] | [bruta] | [salud] | [liquida] |
| [Cia 2] | [uf] | [bruta] | [salud] | [liquida] |
| [Cia 3] | [uf] | [bruta] | [salud] | [liquida] |
| [Cia 4] | [uf] | [bruta] | [salud] | [liquida] |

[INSTRUCCIÓN CRÍTICA: Debes generar AQUI una sección para CADA modalidad de "Renta Vitalicia Inmediata Garantizada" encontrada en el SCOMP (ej. 120 meses, 240 meses). NO LAS OMITAS por ningún motivo. Si hay varias rentas garantizadas, haz una tabla separada para cada una.]

**Renta Vitalicia Inmediata Garantizada [X] Meses** (Repetir para cada periodo encontrado)
**Descripción:** En esta modalidad, si el asegurado fallece durante el periodo garantizaro (ej. [X] meses), los beneficiarios designados recibirán el 100% de la pensión hasta cumplir dicho plazo.
**Cuadro de resultados (4 mejores ofertas):**
| Compañía de Seguros | Pensión en UF | Pensión Bruta $| Descuento 7% Salud$ | Pensión Líquida $ |
| :--- | :--- | :--- | :--- | :--- |
| [Cia 1] | [uf] | [bruta] | [salud] | [liquida] |
| ... | ... | ... | ... | ... |

**Renta Vitalicia Aumentada**
**Descripción:** La "Cláusula de Aumento Temporal de Pensión" es una cobertura adicional que permite duplicar (aumentar en un 100%) el monto de la pensión durante un período determinado al inicio. Una vez que este período finaliza, la pensión vuelve a su monto base original, el cual es fijo en UF y se paga de por vida.
[Generar una sección para CADA modalidad de Renta Vitalicia Aumentada encontrada en el SCOMP, ej: "Renta Vitalicia Aumentada 100% por 36 Meses"]
**[Título de la Modalidad, ej: Renta Vitalicia Aumentada 100% por 36 Meses, Garantizado 180 meses.]**
**Cuadro de resultados (4 mejores ofertas):**
| Compañía | Pensión Aumentada en UF | Pensión Aumentada en $| Descuento 7% Salud$ | Pensión Líquida Período Aumentado | Pensión Después de Aumento en UF (Base) |
| :--- | :--- | :--- | :--- | :--- | :--- |
| [Cia 1] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
| [Cia 2] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
| [Cia 3] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
| [Cia 4] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
**Explicación:** Después del período aumentado, su pensión bajará al monto de la pensión base calculada. En este caso, la mejor oferta es de **[Base UF de la mejor oferta] UF**, lo que equivale a **$[Monto Base $]** brutos.
"""

# === PROMPT PASO 2: RECOMENDACIÓN (SECCIÓN 6) ===
PROMPT_RECOMENDACION = """
Eres un Asesor Previsional experto. Tu tarea es redactar la **Sección 6: Recomendación Final** para un informe.
Te entregaré el análisis de datos (Secciones 1-5) como contexto, y las instrucciones del asesor humano.
Redacta ÚNICAMENTE la "## 6) Recomendación Final" siguiendo las instrucciones.
---
INSTRUCCIONES DEL ASESOR HUMANO PARA LA RECOMENDACIÓN:
"{INSTRUCCIONES_USUARIO}"
---
CONTEXTO (ANÁLISIS DE DATOS SECCIONES 1-5):
{ANALISIS_PREVIO}
---
Redacta ÚNICAMENTE la "## 6) Recomendación Final":
"""

# === PROMPT PASO 3: MODIFICACIÓN ===
PROMPT_MODIFICACION = """
Eres un editor profesional. Tu tarea es tomar el siguiente informe previsional y modificarlo según las instrucciones del usuario.
REGLAS:
1.  **Aplica las modificaciones solicitadas** de forma precisa.
2.  **No cambies el formato Markdown** (títulos ##, ###, tablas |, etc.) a menos que la instrucción te lo pida.
3.  **Mantén el tono profesional** del informe.
4.  Entrega el **informe completo modificado**, no solo la parte que cambiaste.
---
INFORME ORIGINAL:
{INFORME_ACTUAL}
---
INSTRUCCIONES DEL USUARIO PARA MODIFICAR:
"{INSTRUCCIONES_MODIFICACION}"
---
INFORME MODIFICADO:
"""

# === PROMPT PASO 4: VERIFICACIÓN (AUDITORÍA) ===
PROMPT_VERIFICACION = """
Eres un Auditor de Calidad (QC) experto en informes previsionales. Tu misión es revisar que el "Informe Generado" sea fiel a los "Documentos Originales".
NO debes reescribir el informe, solo auditarlo.

Debes verificar DOS cosas críticas:
1.  **Integridad de Modalidades:** ¿Están TODAS las modalidades de pensión del SCOMP en el informe?
    *   **CRÍTICO:** Verifica que se incluyan las **"Renta Vitalicia Inmediata Garantizada"** (ej. 120, 240 meses) si aparecen en el original. Es el error más común.
    *   Si es Invalidez, verifica la "Pensión de Referencia Garantizada".
2.  **Exactitud de Montos:** ¿Los montos en UF de las ofertas coinciden con el documento original?

Documentos Originales (Texto extraído):
{CONTEXTO_ORIGINAL}
---
Informe Generado:
{INFORME_GENERADO}
---
Respuesta del Auditor:
Si todo está correcto y completo, responde EXACTAMENTE: "APROBADO".
Si encuentras errores u omisiones (especialmente modalidades faltantes), responde: "RECHAZADO: [Lista breve de lo que falta o está mal]".
"""

@st.cache_data(show_spinner=False)
def generar_modificacion_ia(informe_actual, instrucciones, api_key):
    """
    Llama a la API de Gemini para MODIFICAR un informe ya existente.
    """
    if not api_key:
        st.error("API Key no configurada.")
        return None
    if not informe_actual or not instrucciones:
        st.error("Faltan datos para modificar el informe.")
        return None
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        prompt_completo = PROMPT_MODIFICACION.format(
            INFORME_ACTUAL=informe_actual,
            INSTRUCCIONES_MODIFICACION=instrucciones
        )
        
        generation_config = {"temperature": 0.2, "response_mime_type": "text/plain"}
        request_options = {"timeout": 300}
        
        response = model.generate_content(
            prompt_completo,
            generation_config=generation_config,
            request_options=request_options
        )
        return response.text
    except Exception as e:
        st.error(f"Error al modificar el informe con IA: {e}")
        st.exception(e)
        return None


@st.cache_data(show_spinner=False)
def generar_analisis_ia(contexto, api_key):
    """
    Llama a la API de Gemini para generar el ANÁLISIS (Secciones 1-5).
    """
    if not api_key:
        st.error("API Key no configurada.")
        return None
    if not contexto:
        st.error("Contexto de PDF vacío.")
        return None
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        fecha_hoy_str = datetime.now().strftime('%d de %B de %Y')
        prompt_completo = PROMPT_ANALISIS.format(
            CONTEXTO_DOCUMENTOS=contexto,
            FECHA_HOY=fecha_hoy_str
        )
        
        generation_config = {"temperature": 0.1, "response_mime_type": "text/plain"}
        request_options = {"timeout": 300}  
        
        response = model.generate_content(
            prompt_completo,
            generation_config=generation_config,
            request_options=request_options
        )
        return response.text
    except Exception as e:
        st.error(f"Error al generar el análisis con IA: {e}")
        st.exception(e)
        return None

@st.cache_data(show_spinner=False)
def generar_recomendacion_ia(analisis_previo, instrucciones, api_key):
    """
    Llama a la API de Gemini para generar SOLO la RECOMENDACIÓN (Sección 6).
    """
    if not api_key:
        st.error("API Key no configurada.")
        return None
    if not analisis_previo or not instrucciones:
        st.error("Faltan datos para generar la recomendación.")
        return None
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        prompt_completo = PROMPT_RECOMENDACION.format(
            ANALISIS_PREVIO=analisis_previo,
            INSTRUCCIONES_USUARIO=instrucciones
        )
        
        generation_config = {"temperature": 0.2, "response_mime_type": "text/plain"}
        request_options = {"timeout": 120}
        
        response = model.generate_content(
            prompt_completo,
            generation_config=generation_config,
            request_options=request_options
        )
        return response.text
    except Exception as e:
        st.error(f"Error al generar la recomendación con IA: {e}")
        st.exception(e)
        return None

@st.cache_data(show_spinner=False)
def verificar_consistencia_ia(contexto, informe_generado, api_key):
    """
    Llama a la API para auditar el informe generado.
    """
    if not api_key: return None
    
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        prompt_completo = PROMPT_VERIFICACION.format(
            CONTEXTO_ORIGINAL=contexto,
            INFORME_GENERADO=informe_generado
        )
        
        generation_config = {"temperature": 0.0, "response_mime_type": "text/plain"}
        
        response = model.generate_content(prompt_completo, generation_config=generation_config)
        return response.text
    except Exception as e:
        # Si falla la auditoría, no bloqueamos el flujo, solo avisamos
        print(f"Error en auditoría IA: {e}")
        return None



# --- 3. FUNCIONES DE DESCARGA (SOLO DOCX) ---

def crear_reporte_doc(informe_texto):
    """
    Crea un archivo .docx en memoria, interpretando Markdown,
    con fuente "Roboto" y sin asteriscos.
    """
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Roboto'
    font.size = Pt(11)

    styles = doc.styles
    for h_level in [1, 2, 3, 4]:
        try:
            h_style = styles[f'Heading {h_level}']
            h_style.font.name = 'Roboto'
            h_style.font.bold = True
        except KeyError:
            pass
            
    try:
        bullet_style = styles['List Bullet']
        bullet_style.font.name = 'Roboto'
        bullet_style.font.size = Pt(11)
    except KeyError:
        pass

    in_table = False
    table = None
    
    for line in informe_texto.split('\n'):
        # Limpia asteriscos para evitar problemas en Word
        line_stripped = line.strip().replace('*', '')

        if line.strip().startswith('|') and line.strip().endswith('|'):
            cells = [c.strip().replace('*', '') for c in line.strip().split('|')[1:-1]]
            
            if '---' in cells[0]:
                continue

            if not in_table:
                try:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, item in enumerate(cells):
                        hdr_cells[i].text = item
                        run = hdr_cells[i].paragraphs[0].runs[0]
                        run.font.name = 'Roboto'
                        run.font.bold = True
                    in_table = True
                except Exception as e:
                    st.warning(f"Error al crear cabecera de tabla DOCX: {e}")
            else:
                try:
                    row_cells = table.add_row().cells
                    for i, item in enumerate(cells):
                         if i < len(row_cells):
                            row_cells[i].text = item
                            run = row_cells[i].paragraphs[0].runs[0]
                            run.font.name = 'Roboto'
                except Exception as e:
                   st.warning(f"Error al añadir fila a tabla DOCX: {e}")
        
        else:
            if in_table:
                doc.add_paragraph()  
                in_table = False
                table = None

            if line.strip().startswith('## '):
                doc.add_heading(line_stripped.replace('## ', ''), level=2)
            elif line.strip().startswith('### '):
                doc.add_heading(line_stripped.replace('### ', ''), level=3)
            elif line.strip().startswith('#### '):
                doc.add_heading(line_stripped.replace('#### ', ''), level=4)
            elif line.strip().startswith('* '):
                doc.add_paragraph(line_stripped, style='List Bullet')
            elif line_stripped and not line_stripped.startswith('---'):
                p = doc.add_paragraph()
                p.add_run(line_stripped)

    if in_table:
        doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# --- 4. LÓGICA PRINCIPAL DE LA APLICACIÓN ---

st.title("🤖 Asistente de Asesoría Previsional (IA)")
st.write("Carga todos los documentos de tu cliente (SCOMP, Cartolas, APV, etc.) para generar un informe de asesoría consolidado.")

# --- Estados de Sesión ---
if 'contexto_documentos' not in st.session_state:
    st.session_state.contexto_documentos = None
# Esta es la ÚNICA variable que guarda el texto del informe
if 'informe_actual' not in st.session_state:
    st.session_state.informe_actual = None
# -------------------------

uploaded_files = st.file_uploader(
    "1. Cargar antecedentes del cliente (PDF)", 
    type=["pdf"],
    accept_multiple_files=True
)

st.divider()

# --- PASO 1: Generar Análisis (Secciones 1-5) ---
if uploaded_files:
    # Esta línea ahora solo se ejecuta si los archivos cambian
    with st.spinner("Leyendo y procesando los archivos PDF..."):
        st.session_state.contexto_documentos = leer_pdfs_cargados(uploaded_files)
    
    if st.button("Generar Análisis de Datos (Secciones 1-5)", type="primary"):
        try:
            final_api_key = st.secrets["api_key"]
        except:
            st.error("Error: La API Key no está configurada en los 'secrets' de la aplicación.")
            final_api_key = None
        
        if final_api_key and st.session_state.contexto_documentos:
            with st.spinner("La IA está analizando los datos (Secciones 1-5)... (Esto puede tardar hasta 1 minuto)"):
                analisis_resultado = generar_analisis_ia(
                    st.session_state.contexto_documentos,
                    final_api_key
                )
            
            if analisis_resultado:
                st.session_state.informe_actual = analisis_resultado # Guarda el análisis (1-5)
                st.success("Análisis (Secciones 1-5) generado.")
                
                # --- AUTO-VERIFICACIÓN ---
                with st.spinner("🔍 El Auditor Virtual está revisando la consistencia del informe..."):
                    resultado_auditoria = verificar_consistencia_ia(
                        st.session_state.contexto_documentos,
                        analisis_resultado,
                        final_api_key
                    )
                
                if resultado_auditoria:
                    if "APROBADO" in resultado_auditoria:
                        st.success("✅ Auditoría Aprobada: El informe incluye todas las modalidades detectadas.")
                    else:
                        st.error("⚠️ Auditoría Detectó Posibles Omisiones:")
                        st.warning(resultado_auditoria)
                        st.info("Revisa si falta alguna modalidad importante (como RV Garantizada). Puedes usar el botón 'Refrescar Informe con Modificaciones' para pedirle a la IA que la agregue.")
                # -------------------------

                st.info("Ya puedes modificar el informe o añadir la recomendación.")
            else:
                st.error("No se pudo generar el análisis.")
        elif not st.session_state.contexto_documentos:
             st.error("Error: No se pudo leer el contexto de los PDF.")

# --- Lógica de Refresco (Sidebar) ---
# Este botón ahora funciona si el informe_actual (Sec 1-5) existe
if st.sidebar.button("Refrescar Informe con Modificaciones"):
    if st.session_state.informe_actual and st.session_state.instrucciones_mod:
        try:
            final_api_key = st.secrets["api_key"]
        except:
            st.error("Error: API Key no configurada.")
            final_api_key = None
        
        if final_api_key:
            with st.spinner("La IA está aplicando tus modificaciones..."):
                informe_modificado = generar_modificacion_ia(
                    st.session_state.informe_actual, # Envía el informe actual (1-5 o 1-6)
                    st.session_state.instrucciones_mod,
                    final_api_key
                )
            if informe_modificado:
                st.session_state.informe_actual = informe_modificado # Sobrescribe el informe
                st.success("Informe refrescado.")
                
                # --- CORRECCIÓN (Limpiar caja de texto) ---
                st.rerun()
                # --- FIN CORRECCIÓN ---
                
            else:
                st.error("No se pudo modificar el informe.")
    elif not st.session_state.informe_actual:
        st.sidebar.warning("Debes generar el 'Análisis de Datos' (Sección 1-5) primero.")
    else:
        st.sidebar.warning("Escribe alguna instrucción de modificación en la caja de texto.")


# --- Botón "Nuevo Informe" ---
st.sidebar.divider()
if st.sidebar.button("Nuevo Informe"):
    # Borra los datos principales del informe y los documentos
    st.session_state.informe_actual = None
    st.session_state.contexto_documentos = None
    
    # Limpia las cajas de texto
    if 'instrucciones_rec' in st.session_state:
        st.session_state.instrucciones_rec = ""
    
    st.success("Memoria limpiada. Puedes cargar un nuevo informe.")
    st.rerun()
# --- FIN Botón "Nuevo Informe" ---

# --- Botón "Generar Contrato" (VERSIÓN DOCX + FORMULARIO) ---
st.sidebar.divider()
st.sidebar.subheader("Generar Contrato (DOCX)")
tipo_contrato_sel = st.sidebar.selectbox("Tipo de Contrato", ["Vejez o Invalidez", "Sobrevivencia"])

# Lógica del Formulario
if st.session_state.informe_actual:
    # 1. Intentar extraer datos automáticamente del informe
    if 'contract_data' not in st.session_state or st.sidebar.button("Recargar Datos de Informe"):
        st.session_state.contract_data = extract_contract_data(st.session_state.informe_actual)
    
    st.sidebar.write("Completa los datos para el contrato:")
    
    with st.sidebar.form("contract_form"):
        # Campos de texto editables (pre-llenados con lo extraído)
        # Nombres de variables coincidentes con lo que BUSCAREMOS en el DOCX para reemplazar
        # Si el DOCX no tiene placeholders, buscaremos estos textos literales o los insertaremos.
        # Definiremos un estándar aquí: Reemplazar "____________________" si es un campo vacío, o usar placeholders si el usuario los pone.
        # Dado que no sabemos, pediremos input y trataremos de reemplazar.
        
        c_nombre = st.text_input("Nombre Completo", value=st.session_state.contract_data.get("Nombre Completo", ""))
        c_rut = st.text_input("RUT", value=st.session_state.contract_data.get("RUT", ""))
        c_direccion = st.text_input("Dirección", value=st.session_state.contract_data.get("Dirección", ""))
        c_comuna = st.text_input("Comuna", value=st.session_state.contract_data.get("Comuna", ""))
        c_ciudad = st.text_input("Ciudad", value=st.session_state.contract_data.get("Ciudad", st.session_state.contract_data.get("Comuna", "")))
        c_telefono = st.text_input("Teléfono", value=st.session_state.contract_data.get("Teléfono", ""))
        c_celular = st.text_input("Celular", value=st.session_state.contract_data.get("Celular", ""))
        c_email = st.text_input("Correo Electrónico", value=st.session_state.contract_data.get("Correo Electrónico", ""))
        c_estado_civil = st.text_input("Estado Civil", value=st.session_state.contract_data.get("Estado Civil", ""))
        c_fecha_nac = st.text_input("Fecha de Nacimiento", value=st.session_state.contract_data.get("Fecha de Nacimiento", ""))
        c_oficio = st.text_input("Profesión u Oficio", value=st.session_state.contract_data.get("Profesión u Oficio", ""))
        c_afp = st.text_input("AFP de Origen", value=st.session_state.contract_data.get("AFP de Origen", ""))
        c_salud = st.text_input("Sistema de Salud", value=st.session_state.contract_data.get("Sistema de Salud", st.session_state.contract_data.get("Institución de Salud", "")))
        c_tipo_pension = st.text_input("Tipo de Pensión", value=st.session_state.contract_data.get("Tipo de Pensión Solicitada", ""))
        c_fecha = st.text_input("Fecha", value=datetime.now().strftime("%d/%m/%Y"))
        
        # --- Campos Adicionales para Sobrevivencia ---
        c_causante_nombre = ""
        c_causante_rut = ""
        c_consultante_nombre = ""
        c_consultante_rut = ""
        
        c_ben_nombre = ""
        c_ben_rut = ""
        c_ben_parentesco = ""
        c_ben_sexo = "" # Nuevo
        c_ben_invalidez = "" # Nuevo
        c_ben_nac = "" # Nuevo
        
        if tipo_contrato_sel == "Sobrevivencia":
            st.markdown("---")
            st.markdown("### Datos de Sobrevivencia")
            
            st.caption("Causante (Fallecido)")
            col1, col2 = st.columns(2)
            with col1:
                c_causante_nombre = st.text_input("Nombre Causante", value=st.session_state.contract_data.get("Causante Nombre", ""))
            with col2:
                c_causante_rut = st.text_input("RUT Causante", value=st.session_state.contract_data.get("Causante RUT", ""))
                
            st.caption("Consultante (Solicitante)")
            col3, col4 = st.columns(2)
            with col3:
                c_consultante_nombre = st.text_input("Nombre Consultante", value=st.session_state.contract_data.get("Consultante Nombre", ""))
            with col4:
                c_consultante_rut = st.text_input("RUT Consultante", value=st.session_state.contract_data.get("Consultante RUT", ""))

            st.markdown("### Datos del Beneficiario Principal")
            st.info("El sistema precarga el beneficiario 1. Si hay más, edita el contrato final.")
            
            c_ben_nombre = st.text_input("Nombre Beneficiario", value=st.session_state.contract_data.get("Beneficiario 1 Nombre", ""))
            c_ben_rut = st.text_input("RUT Beneficiario", value=st.session_state.contract_data.get("Beneficiario 1 RUT", ""))
            c_ben_parentesco = st.text_input("Parentesco", value=st.session_state.contract_data.get("Beneficiario 1 Parentesco", ""))
            
            col_b1, col_b2, col_b3 = st.columns(3)
            with col_b1:
                 c_ben_sexo = st.text_input("Sexo", value=st.session_state.contract_data.get("Beneficiario 1 Sexo", ""))
            with col_b2:
                 c_ben_invalidez = st.text_input("Invalidez", value=st.session_state.contract_data.get("Beneficiario 1 Invalidez", ""))
            with col_b3:
                 c_ben_nac = st.text_input("Fecha Nac.", value=st.session_state.contract_data.get("Beneficiario 1 Fecha de Nacimiento", "")) # Nuevo
        # ---------------------------------------------
        
        submitted = st.form_submit_button("Generar DOCX")
        
        if submitted:
            # Preparar diccionario de reemplazos
            # Mapeo EXACTO a los placeholders del DOCX (según screenshot del usuario)
            replacements = {
                "{NOMBRE AFILIADO}": c_nombre,
                "{RUT AFILIADO}": c_rut,
                "{DIRECCIÓN}": c_direccion,
                "{COMUNA}": c_comuna,
                "{CIUDAD}": c_ciudad,
                "{CELULAR}": c_celular if c_celular else c_telefono, # Priorizar celular si existe
                "{TELEFONO}": c_telefono,
                "{CORREO ELECTRÓNICO}": c_email,
                "{ESTADO CIVIL AFILIADO}": c_estado_civil,
                "{FECHA DE NACIMIENTO AFILIADO}": c_fecha_nac,
                "{OFICIO AFILIADO}": c_oficio,
                "{AFP DE ORIGEN}": c_afp,
                "{SISTEMA DE SALUD}": c_salud,
                "{TIPO DE PENSIÓN}": c_tipo_pension,
                "{FECHA}": c_fecha,
                # Compatibilidad con posibles otros placeholders
                "{{RUT}}": c_rut,
            }
            
            # --- Agregado para Sobrevivencia ---
            if tipo_contrato_sel == "Sobrevivencia":
                replacements["{NOMBRE CAUSANTE}"] = c_causante_nombre
                replacements["{RUT CAUSANTE}"] = c_causante_rut
                replacements["{NOMBRE CONSULTANTE}"] = c_consultante_nombre
                replacements["{RUT CONSULTANTE}"] = c_consultante_rut
                
                replacements["{NOMBRE BENEFICIARIO}"] = c_ben_nombre
                replacements["{RUT BENEFICIARIO}"] = c_ben_rut
                replacements["{PARENTESCO BENEFICIARIO}"] = c_ben_parentesco
                replacements["{SEXO BENEFICIARIO}"] = c_ben_sexo
                replacements["{INVALIDEZ BENEFICIARIO}"] = c_ben_invalidez
                replacements["{FECHA NAC BENEFICIARIO}"] = c_ben_nac
                # Si el contrato tiene placeholders para mas beneficiarios, se pueden agregar aqui
                # Si el contrato tiene placeholders para mas beneficiarios, se pueden agregar aqui
                # Por ahora asumimos 1 principal o que el usuario rellena el resto a mano si son muchos.
            # -----------------------------------
            
            # --- NUEVA LÓGICA: Recopilar TODOS los beneficiarios para la tabla ---
            beneficiaries_list = []
            if tipo_contrato_sel == "Sobrevivencia":
                # Recorrer todos los beneficiarios encontrados en st.session_state.contract_data
                # Buscamos Beneficiario 1, Beneficiario 2, etc.
                for i in range(1, 10): # Hasta 10 por seguridad
                    if f"Beneficiario {i} Nombre" in st.session_state.contract_data:
                        # Crear diccionario para ESTE beneficiario
                        ben_dict = {
                            "{NOMBRE BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} Nombre", ""),
                            "{RUT BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} RUT", ""),
                            "{PARENTESCO BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} Parentesco", ""),
                            "{SEXO BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} Sexo", ""),
                            "{INVALIDEZ BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} Invalidez", ""),
                            "{FECHA NAC BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} Fecha de Nacimiento", ""),
                             # Variant
                            "{FECHA NACIMIENTO BENEFICIARIO}": st.session_state.contract_data.get(f"Beneficiario {i} Fecha de Nacimiento", ""),
                        }
                        # Opción de sobreescribir el primero con los datos del form (si el usuario editó el 1)
                        if i == 1:
                            ben_dict["{NOMBRE BENEFICIARIO}"] = c_ben_nombre
                            ben_dict["{RUT BENEFICIARIO}"] = c_ben_rut
                            ben_dict["{PARENTESCO BENEFICIARIO}"] = c_ben_parentesco
                            ben_dict["{SEXO BENEFICIARIO}"] = c_ben_sexo
                            ben_dict["{INVALIDEZ BENEFICIARIO}"] = c_ben_invalidez
                            ben_dict["{FECHA NAC BENEFICIARIO}"] = c_ben_nac
                            ben_dict["{FECHA NACIMIENTO BENEFICIARIO}"] = c_ben_nac
                        
                        beneficiaries_list.append(ben_dict)
                    else:
                        # Si es el 1 y no estaba en data (caso borde manual), lo agregamos del form
                        if i == 1 and c_ben_nombre:
                             ben_dict = {
                                "{NOMBRE BENEFICIARIO}": c_ben_nombre,
                                "{RUT BENEFICIARIO}": c_ben_rut,
                                "{PARENTESCO BENEFICIARIO}": c_ben_parentesco,
                                "{SEXO BENEFICIARIO}": c_ben_sexo,
                                "{INVALIDEZ BENEFICIARIO}": c_ben_invalidez,
                                "{FECHA NAC BENEFICIARIO}": c_ben_nac,
                                "{FECHA NACIMIENTO BENEFICIARIO}": c_ben_nac,
                            }
                             beneficiaries_list.append(ben_dict)

            # ---------------------------------------------------------------------

            template_file = get_contract_template_path(tipo_contrato_sel)
            
            with st.spinner("Generando contrato..."):
                try:
                    docx_bytes = generate_contract_docx(template_file, replacements, beneficiaries_list)
                    
                    st.session_state['ultimo_contrato_docx'] = docx_bytes
                    st.session_state['ultimo_contrato_name'] = f"Contrato_Final_{c_nombre.split()[0]}.docx"
                    st.success("¡Contrato Generado!")
                except Exception as e:
                    st.error(f"Error: {e}")

else:
    st.sidebar.info("Genera el análisis primero para precargar datos.")

if 'ultimo_contrato_docx' in st.session_state:
    st.sidebar.download_button(
        label="⬇️ Descargar Contrato Final",
        data=st.session_state['ultimo_contrato_docx'],
        file_name=st.session_state['ultimo_contrato_name'],
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
# --- FIN Botón "Generar Contrato" ---


# --- PASO 2 y 3: Mostrar Informe, Pedir Recomendación y Descargar ---
# Esta sección ahora se muestra solo si el informe_actual existe.
if st.session_state.informe_actual:
    
    st.subheader("Vista Previa del Informe Actual")
    st.markdown(st.session_state.informe_actual)
    
    st.divider()
    st.subheader("2. Instrucciones para la Recomendación Final (Sección 6)")
    # Este widget se dibuja aquí
    st.text_area(
        "Escriba sus instrucciones para la recomendación:", 
        key="instrucciones_rec", 
        height=150,
        help="Escribe aquí tus ideas (ej. 'Recomendar RVA a 60m...') y presiona el botón de abajo para AÑADIR la Sección 6 al informe."
    )

    if st.button("Añadir Recomendación al Informe (Sección 6)", type="primary"):
        
        try:
            final_api_key = st.secrets["api_key"]
        except:
            st.error("Error: La API Key no está configurada en los 'secrets' de la aplicación.")
            final_api_key = None

        instrucciones_texto = st.session_state.instrucciones_rec
        
        if final_api_key and instrucciones_texto:
            with st.spinner("La IA está redactando y añadiendo la recomendación (Sección 6)..."):
                recomendacion_resultado = generar_recomendacion_ia(
                    st.session_state.informe_actual, # Usa el informe actual (1-5) como contexto
                    instrucciones_texto,
                    final_api_key
                )
            
            if recomendacion_resultado:
                # --- Lógica de AÑADIR ---
                st.session_state.informe_actual += "\n\n" + recomendacion_resultado
                
                # --- CORRECCIÓN (Limpiar caja de texto) ---
                
                st.success("Recomendación añadida. Ya puedes modificar el informe completo o descargarlo.")
                
                # Refresca la página
                st.rerun() 
                # --- FIN CORRECCIÓN ---
                
            else:
                st.error("No se pudo generar la recomendación.")
        elif not instrucciones_texto:
            st.warning("Por favor, escriba las instrucciones para la recomendación.")

    # --- Sección de Descarga ---
    st.divider()
    st.subheader("Descargar Informe Completo")
    
    try:
        informe_completo_texto = st.session_state.informe_actual
        
        doc_data = crear_reporte_doc(informe_completo_texto)
        
        # Extraer nombre para el archivo
        nombre_cliente = "Cliente"
        match = re.search(r'\*\*Nombre Completo:\*\*\s*(.*)', informe_completo_texto)
        if match:
            nombre_cliente = match.group(1).strip().replace("/", "-").replace("\\", "-")
        
        file_name_download = f"Informe_final_Asesoria_Previsional_{nombre_cliente}.docx"

        st.download_button(
            label="📄 Descargar Informe en DOCX (Compatible con Word/Google Docs)",
            data=doc_data,
            file_name=file_name_download,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
            
    except Exception as e:
        st.error(f"Error al generar el archivo de descarga: {e}")
        st.exception(e)

# No debe haber ninguna línea "st.session_state.instrucciones_..." aquí al final.
