import re
import io
import logging
from typing import Dict, Optional, Any
from pathlib import Path
from docx import Document  # type: ignore
from docx.text.paragraph import Paragraph  # type: ignore

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- CONSTANTS ---
# Using absolute paths or relative to the script location is safer
BASE_DIR = Path(__file__).parent.resolve()
TEMPLATE_OLD_AGE_FILENAME = "CONTRATO 2026 AP - PLANTILLA.docx"
TEMPLATE_SURVIVORSHIP_FILENAME = "CONTRATO 2026 sobrevivencia -  plantilla.docx"

def get_contract_template_path(contract_type: str) -> Path:
    """
    Returns the path to the DOCX template based on the contract type.
    
    Args:
        contract_type: The type of contract (e.g., "Vejez o Invalidez").
        
    Returns:
        Path object to the template file.
        
    Raises:
        FileNotFoundError: If the template file does not exist.
    """
    if contract_type == "Vejez o Invalidez":
        filename = TEMPLATE_OLD_AGE_FILENAME
    else:
        filename = TEMPLATE_SURVIVORSHIP_FILENAME
        
    template_path = BASE_DIR / filename
    
    if not template_path.exists():
        logger.error(f"Template not found at: {template_path}")
        raise FileNotFoundError(f"Template file not found: {filename}")
        
    return template_path

def extract_contract_data(markdown_text: str) -> Dict[str, str]:
    """
    Extracts relevant contract data (Name, RUT, Address, etc.) from a Markdown report.
    
    Args:
        markdown_text: The content of the markdown report.
        
    Returns:
        A dictionary with extracted fields.
    """
    data: Dict[str, str] = {}
    if not markdown_text:
        return data

    # Regex patterns for extraction
    # Using case-insensitive search and handling potential markdown formatting variations
    patterns = {
        "Nombre Completo": r"\*\*Nombre Completo:?\*\*\s*(.*)",
        "RUT": r"\*\*RUT:?\*\*\s*(.*)",
        # Updated to catch "Dirección", "Direccion", "Domicilio", etc.
        "Dirección": r"\*\*(?:Direcci[óo]n|Domicilio):?\*\*\s*(.*)", 
        "Comuna": r"\*\*Comuna:?\*\*\s*(.*)",
        "Ciudad": r"\*\*Ciudad:?\*\*\s*(.*)",
        "Teléfono": r"\*\*Tel[ée]fono:?\*\*\s*(.*)",
        "Celular": r"\*\*Celular:?\*\*\s*(.*)",
        "Correo Electrónico": r"\*\*Correo.*?:?\*\*\s*(.*)",
        "Estado Civil": r"\*\*Estado Civil:?\*\*\s*(.*)",
        "Cédula de Identidad": r"\*\*Cédula.*?:?\*\*\s*(.*)",
        "Fecha de Nacimiento": r"\*\*Fecha de Nacimiento:?\*\*\s*(.*)",
        "AFP de Origen": r"\*\*AFP de Origen:?\*\*\s*(.*)",
        "Institución de Salud": r"\*\*Institución de Salud:?\*\*\s*(.*)",
        "Sistema de Salud": r"\*\*Sistema de Salud:?\*\*\s*(.*)",
        "Tipo de Pensión Solicitada": r"\*\*Tipo de Pensión Solicitada:?\*\*\s*(.*)",
        # New Fields for Solicitud de Ofertas
        "Fecha Solicitud de Ofertas": r"\*\*Fecha Solicitud de Ofertas:?\*\*\s*(.*)",
        "Modalidades Solicitadas": r"\*\*Modalidades Solicitadas:?\*\*\s*(.*)",
        # Survivorship Specifics
        "Causante Nombre": r"\*\*Causante Nombre:?\*\*\s*(.*)",
        "Causante RUT": r"\*\*Causante RUT:?\*\*\s*(.*)",
        "Consultante Nombre": r"\*\*Consultante Nombre:?\*\*\s*(.*)",
        "Consultante RUT": r"\*\*Consultante RUT:?\*\*\s*(.*)",
    }
    
    for field, pattern in patterns.items():
        match = re.search(pattern, markdown_text, re.IGNORECASE)
        if match:
            # Clean up the extracted text (remove trailing periods, markdown artifacts)
            val = match.group(1).strip()
            # Filter out placeholder or invalid values
            if val and "No informada" not in val and "[Extraer" not in val:
                data[field] = val
            
    # Fallback for RUT if not found by specific label
    if "RUT" not in data:
         # Looks for standard Chilean RUT format: XX.XXX.XXX-X
         rut_match = re.search(r"\b(\d{1,2}\.\d{3}\.\d{3}-[\dkK])\b", markdown_text)
         if rut_match:
             data["RUT"] = rut_match.group(1)

    # --- Beneficiary Extraction (Table Parsing) ---
    # Looks for the table in Section 2) Antecedentes del beneficiario
    # We parse the markdown table to extract beneficiaries 1 to 5
    
    # 1. Locate Section 2
    section_2_match = re.search(r"### 2\) Antecedentes del beneficiario.*?(?=### 3\))", markdown_text, re.DOTALL)
    if section_2_match:
        section_2_text = section_2_match.group(0)
        
        # 2. Find table rows (text between pipes |)
        # Skip header and separator lines
        lines = section_2_text.split('\n')
        ben_index = 1
        
        for line in lines:
            if "|" in line and "Nombre Completo" not in line and "---" not in line:
                # This is likely a data row
                cols = [c.strip() for c in line.split('|') if c.strip()]
                
                # We expect at least these columns: Name, RUT, Parentesco...
                # | Nombre Completo | RUT | Parentesco | Sexo | Invalidez | Fecha de Nacimiento |
                if len(cols) >= 3:
                     data[f"Beneficiario {ben_index} Nombre"] = cols[0]
                     data[f"Beneficiario {ben_index} RUT"] = cols[1]
                     data[f"Beneficiario {ben_index} Parentesco"] = cols[2]
                     
                     if len(cols) >= 4: data[f"Beneficiario {ben_index} Sexo"] = cols[3]
                     if len(cols) >= 5: data[f"Beneficiario {ben_index} Invalidez"] = cols[4]
                     if len(cols) >= 6: data[f"Beneficiario {ben_index} Fecha de Nacimiento"] = cols[5]
                     
                     ben_index += 1
                     if ben_index > 5: break # Limit to 5
    
    return data

def _replace_text_in_paragraph(paragraph: Paragraph, replacements: Dict[str, str]) -> None:
    """
    Helper function to replace text within a DOCX paragraph.
    Handles both {{KEY}} placeholders and underscore lines (e.g., "Name: _______").
    
    Args:
        paragraph: The paragraph object from python-docx.
        replacements: Dictionary of keys to replace.
    """
    text = paragraph.text
    original_text = text
    
    # 1. Direct Placeholder Replacement {{KEY}}
    for key, value in replacements.items():
        if key in text and value:
            text = text.replace(key, str(value))
            
    # 2. Underscore Line Replacement
    # Example: "Nombre: ________________" -> "Nombre: Juan Perez"
    if "_" * 5 in text:
        # Map labels in the doc to keys in our data dictionary
        # This mapping covers common labels found in the templates
        label_map = {
            "Nombre": replacements.get("{{NOMBRE}}", ""),
            "Señor": replacements.get("{{NOMBRE}}", ""),
            "RUT": replacements.get("{{RUT}}", ""),
            "Dirección": replacements.get("{{DIRECCION}}", ""),
            "Domicilio": replacements.get("{{DIRECCION}}", ""),
            "Comuna": replacements.get("{{COMUNA}}", ""),
            "Teléfono": replacements.get("{{TELEFONO}}", ""),
            "Fecha": replacements.get("{{FECHA}}", ""),
        }
        
        for label, val in label_map.items():
            if val and label in text:
                # Regex to find label followed by dots or underscores
                # e.g., "Nombre: ..........." or "Nombre: ________"
                # using a raw string for regex pattern
                pattern = rf"({label}.*?)([_\\.]+)"
                regex = re.compile(pattern, re.IGNORECASE)
                # Replace with captured label + value
                text = regex.sub(rf"\1 {val}", text)

    # Only update paragraph text if changes were made
    if text != original_text:
        paragraph.text = text

def _fill_beneficiary_table(doc: Document, beneficiaries: list) -> None:
    """
    Fills the beneficiary table in the document with the provided data.
    """
    if not beneficiaries:
        return

    # Identify the correct table. We look for a table that has "Parentesco" or "Nombre" in header
    target_table = None
    header_row_index = 0
    
    for table in doc.tables:
        if not table.rows: continue
        # Check first few rows for headers
        for i, row in enumerate(table.rows[:3]):
            row_text = " ".join([cell.text for cell in row.cells]).lower()
            if "parentesco" in row_text and "rut" in row_text:
                target_table = table
                header_row_index = i
                break
        if target_table:
            break
            
    if not target_table:
        logger.warning("Beneficiary table not found in DOCX template.")
        return

    # Start filling from the row after header
    data_row_start = header_row_index + 1
    
    # We will iterate through the table rows available for data
    # If there are more beneficiaries than rows, we can't add them easily without breaking format,
    # so we'll fill what we can. If fewer, we clear the remaining placeholders.
    
    available_rows = target_table.rows[data_row_start:]
    
    for i, row in enumerate(available_rows):
        # Determine if we have a beneficiary for this row
        if i < len(beneficiaries):
            ben_data = beneficiaries[i]
            # Replace placeholders in this row with specific beneficiary data
            # expected keys in ben_data: "{NOMBRE BENEFICIARIO}", "{RUT BENEFICIARIO}", etc.
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, ben_data)
        else:
            # No beneficiary for this row, clear the placeholders
            # We treat this as "filling with empty strings" for the specific placeholders
            empty_data = {
                "{NOMBRE BENEFICIARIO}": "",
                "{RUT BENEFICIARIO}": "",
                "{PARENTESCO BENEFICIARIO}": "",
                "{SEXO BENEFICIARIO}": "",
                "{INVALIDEZ BENEFICIARIO}": "",
                "{FECHA NAC BENEFICIARIO}": "",
                "{FECHA NACIMIENTO BENEFICIARIO}": "", # Variant
            }
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, empty_data)


def generate_contract_docx(template_path: Path, data: Dict[str, str], beneficiaries_list: Optional[list] = None) -> bytes:
    """
    Generates a filled DOCX contract based on a template and data.
    
    Args:
        template_path: Path to the .docx template file.
        data: Dictionary of data to insert. Keys should match placeholders (e.g., "{{NOMBRE}}").
        beneficiaries_list: Optional list of dictionaries, one for each beneficiary.
        
    Returns:
        Bytes of the generated DOCX file.
        
    Raises:
        Exception: Captures and re-raises any errors during processing.
    """
    try:
        if not template_path.exists():
             raise FileNotFoundError(f"Template not found at {template_path}")

        doc = Document(str(template_path))
        
        # 1. Fill Global Placeholders (Main Contract Data)
        # Process all paragraphs in the document body
        for paragraph in doc.paragraphs:
            _replace_text_in_paragraph(paragraph, data)
            
        # Process all tables (common in forms)
        # Note: We do this for ALL tables to catch global placeholders (like Affiliate Name) inside tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                         _replace_text_in_paragraph(paragraph, data)
        
        # 2. Fill Beneficiary Table (Specific Logic)
        if beneficiaries_list:
            _fill_beneficiary_table(doc, beneficiaries_list)
                        
        # Save to memory buffer
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.getvalue()
        
    except Exception as e:
        logger.error(f"Error generating contract: {e}", exc_info=True)
        raise e
