from docx import Document

def inspect_docx(filename):
    print(f"--- Inspecting: {filename} ---")
    try:
        doc = Document(filename)
        print("--- Paragraphs (First 20) ---")
        for i, p in enumerate(doc.paragraphs[:20]):
            if p.text.strip():
                print(f"P{i}: {repr(p.text)}")
        
        print("\n--- Tables (First 5) ---")
        for i, table in enumerate(doc.tables[:5]):
            print(f"Table {i}:")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    print(f"  Row: {row_text}")
    except Exception as e:
        print(f"Error: {e}")

inspect_docx("CONTRATO 2026 AP - PLANTILLA.docx")
print("\n" + "="*30 + "\n")
inspect_docx("CONTRATO 2026 sobrevivencia -  plantilla.docx")
