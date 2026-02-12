import re
import io
import streamlit as st
from docx import Document
from datetime import datetime

# --- CONSTANTES DE PLANTILLAS DOCX ---
TEMPLATE_VEJEZ_DOCX = "CONTRATO 2026 AP - PLANTILLA.docx" 
TEMPLATE_SOBREVIVENCIA_DOCX = "CONTRATO 2026 sobrevivencia -  plantilla.docx"

def get_template_path(tipo):
    if tipo == "Vejez o Invalidez":
        return TEMPLATE_VEJEZ_DOCX
    else:
        return TEMPLATE_SOBREVIVENCIA_DOCX

def extract_client_data_from_markdown(markdown_text):
    """
    Intenta extraer datos básicos del informe Markdown usando Regex.
    Retorna un diccionario con lo encontrado.
    """
    data = {}
    if not markdown_text:
        return data

    # Patrones de búsqueda comunes en el informe
    # Se modifican para ser más permisivos (ignorar mayúsculas, espacios extra)
    patterns = {
        "Nombre Completo": r"\*\*Nombre Completo:?\*\*\s*(.*)",
        "RUT": r"\*\*RUT:?\*\*\s*(.*)",
        "Dirección": r"\*\*Direcci[óo]n:?\*\*\s*(.*)", 
        "Comuna": r"\*\*Comuna:?\*\*\s*(.*)",
        "Teléfono": r"\*\*Tel[ée]fono:?\*\*\s*(.*)",
        "Estado Civil": r"\*\*Estado Civil:?\*\*\s*(.*)",
        "Cédula de Identidad": r"\*\*Cédula.*:?\*\*\s*(.*)", # A veces cambia
    }
    
    for field, pattern in patterns.items():
        match = re.search(pattern, markdown_text, re.IGNORECASE)
        if match:
            val = match.group(1).strip()
            # Filtrar valores no útiles
            if val and "No informada" not in val and "[Extraer" not in val:
                data[field] = val
            
    # Fallback: Si RUT no se encuentra, buscar patrón de RUT chileno simple
    if "RUT" not in data:
         rut_match = re.search(r"\b(\d{1,2}\.\d{3}\.\d{3}-[\dkK])\b", markdown_text)
         if rut_match:
             data["RUT"] = rut_match.group(1)

    return data

def replace_text_in_paragraph(paragraph, replacements):
    """
    Reemplaza texto en un párrafo.
    Estrategia Mixta:
    1. Reemplazo directo de Placeholders {{KEY}} si existen.
    2. Reemplazo de líneas de llenado (____________________) si se detecta la etiqueta clave cerca.
    """
    text = paragraph.text
    original_text = text
    
    # 1. Reemplazo de Placeholders explícitos {{KEY}}
    # Iteramos sobre las claves para ver si existen en el texto
    for key, value in replacements.items():
        if key in text and value:
            text = text.replace(key, str(value))
            
    # 2. Estrategia de "Líneas de Llenado" (Underscores)
    # Si tenemos "Nombre: ________________", lo convertimos a "Nombre: Juan Perez"
    # Detectamos underscores largos
    if "_" * 5 in text:
        # Mapeo de Etiqueta -> Valor que debería ir ahí
        label_map = {
            "Nombre": replacements.get("{{NOMBRE}}", ""),
            "Señor": replacements.get("{{NOMBRE}}", ""), # A veces dice "Señor(a): ____"
            "RUT": replacements.get("{{RUT}}", ""),
            "Dirección": replacements.get("{{DIRECCION}}", ""),
            "Domicilio": replacements.get("{{DIRECCION}}", ""),
            "Comuna": replacements.get("{{COMUNA}}", ""),
            "Teléfono": replacements.get("{{TELEFONO}}", ""),
            "Fecha": replacements.get("{{FECHA}}", ""),
        }
        
        for label, val in label_map.items():
            if val and label in text:
                # Regex: Busca la etiqueta, seguida de caracteres opcionales y luego la línea
                # Ej: "Nombre: ..........." o "Nombre: ________"
                # Reemplazamos la línea por el VALOR en Negrita (idealmente, pero txt plano aquí)
                regex = re.compile(f"({label}.*?)([_\\.]+)", re.IGNORECASE)
                text = regex.sub(f"\\1 {val}", text)

    # Si hubo cambios, actualizamos el párrafo
    # NOTA: Esto reemplaza todo el estilo del párrafo con el del primer run.
    # Para contratos simples suele ser aceptable.
    if text != original_text:
        paragraph.text = text


def fill_contract_template(template_path, data_dict):
    """
    Abre el DOCX, y realiza reemplazos en parrafos y tablas.
    """
    try:
        doc = Document(template_path)
        
        # 1. Párrafos principales (Cuerpo del contrato)
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, data_dict)
            
        # 2. Tablas (Muy común en formularios)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                         replace_text_in_paragraph(paragraph, data_dict)
                        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
        
    except Exception as e:
        return None, str(e)


