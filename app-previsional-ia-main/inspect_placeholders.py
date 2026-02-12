from docx import Document
import re

def inspect(filename):
    print(f"--- Inspecting {filename} ---")
    try:
        doc = Document(filename)
        text = []
        for p in doc.paragraphs:
            text.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        text.append(p.text)
        
        full_text = "\n".join(text)
        # Find everything looking like {KEY} or {{KEY}}
        matches = re.findall(r"\{+[A-Z0-9 _\-ÑÁÉÍÓÚ]+\}+", full_text)
        print("Found placeholders:", set(matches))
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    inspect("CONTRATO_FIXED_OLD_AGE.docx")
    inspect("CONTRATO_FIXED_SURVIVORSHIP.docx")
