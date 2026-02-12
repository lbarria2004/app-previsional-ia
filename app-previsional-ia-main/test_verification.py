
import unittest
import os
from pathlib import Path
from docx import Document
from contract_utils import extract_contract_data, _fill_beneficiary_table, _replace_text_in_paragraph

class TestBeneficiaryLogic(unittest.TestCase):

    def test_extraction_from_markdown_table(self):
        print("\n--- Testing Markdown Extraction ---")
        markdown_sample = """
### 2) Antecedentes del beneficiario
El afiliado declara a la siguiente beneficiaria legal de pensión:
| Nombre Completo | RUT | Parentesco | Sexo | Invalidez | Fecha de Nacimiento |
| :--- | :--- | :--- | :--- | :--- | :--- |
| Juan Perez | 11.111.111-1 | Hijo | M | N | 01/01/2000 |
| Maria Gomez | 22.222.222-2 | Conyuge | F | N | 02/02/1980 |

### 3) Situación previsional
"""
        data = extract_contract_data(markdown_sample)
        
        # Check Beneficiary 1
        self.assertEqual(data.get("Beneficiario 1 Nombre"), "Juan Perez")
        self.assertEqual(data.get("Beneficiario 1 RUT"), "11.111.111-1")
        self.assertEqual(data.get("Beneficiario 1 Parentesco"), "Hijo")
        self.assertEqual(data.get("Beneficiario 1 Sexo"), "M")
        
        # Check Beneficiary 2
        self.assertEqual(data.get("Beneficiario 2 Nombre"), "Maria Gomez")
        self.assertEqual(data.get("Beneficiario 2 RUT"), "22.222.222-2")
        self.assertEqual(data.get("Beneficiario 2 Parentesco"), "Conyuge")
        
        print("✅ Extraction Verified!")

    def test_docx_table_filling(self):
        print("\n--- Testing DOCX Table Filling ---")
        # 1. Create a dummy DOCX with a table
        doc = Document()
        doc.add_paragraph("Contrato de Prueba")
        
        # Create a table with headers and 3 data rows
        table = doc.add_table(rows=4, cols=3)
        
        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Nombre"
        hdr_cells[1].text = "RUT"
        hdr_cells[2].text = "Parentesco"
        
        # Data Rows with Placeholders
        for i in range(1, 4): # Rows 1, 2, 3
            row_cells = table.rows[i].cells
            row_cells[0].text = "{NOMBRE BENEFICIARIO}"
            row_cells[1].text = "{RUT BENEFICIARIO}"
            row_cells[2].text = "{PARENTESCO BENEFICIARIO}"
            
        # 2. Define Beneficiaries Data
        beneficiaries_list = [
            {
                "{NOMBRE BENEFICIARIO}": "BENEFICIARIO_UNO",
                "{RUT BENEFICIARIO}": "1-9",
                "{PARENTESCO BENEFICIARIO}": "MADRE"
            },
            {
                "{NOMBRE BENEFICIARIO}": "BENEFICIARIO_DOS",
                "{RUT BENEFICIARIO}": "2-9",
                "{PARENTESCO BENEFICIARIO}": "PADRE"
            }
        ]
        
        # 3. Run the filling function
        _fill_beneficiary_table(doc, beneficiaries_list)
        
        # 4. Assertions
        # Row 1 (Index 1) should be Ben 1
        row1_text = " ".join([c.text for c in table.rows[1].cells])
        self.assertIn("BENEFICIARIO_UNO", row1_text)
        self.assertIn("MADRE", row1_text)
        
        # Row 2 (Index 2) should be Ben 2
        row2_text = " ".join([c.text for c in table.rows[2].cells])
        self.assertIn("BENEFICIARIO_DOS", row2_text)
        self.assertIn("PADRE", row2_text)
        
        # Row 3 (Index 3) should be EMPTY (Clearing logic)
        row3_text = " ".join([c.text for c in table.rows[3].cells])
        self.assertNotIn("{NOMBRE BENEFICIARIO}", row3_text) # Placeholder should be gone
        self.assertNotIn("BENEFICIARIO", row3_text) # Should be empty
        self.assertEqual(row3_text.strip(), "")
        
        print("✅ DOCX Table Filling Verified!")
        
        # Optional: Save for manual inspection if needed
        # doc.save("test_output.docx")

if __name__ == '__main__':
    with open('test_results.txt', 'w', encoding='utf-8') as f:
        runner = unittest.TextTestRunner(stream=f, verbosity=2)
        unittest.main(testRunner=runner, exit=False)

