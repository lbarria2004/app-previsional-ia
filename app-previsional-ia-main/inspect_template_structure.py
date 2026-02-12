
from docx import Document
import os

def inspect_template(filename):
    if not os.path.exists(filename):
        print(f"File not found: {filename}")
        return

    doc = Document(filename)
    print(f"Inspecting: {filename}")
    print("-" * 30)

    print(f"Total Tables: {len(doc.tables)}")
    
    print("\n--- Searching for {NOMBRE BENEFICIARIO} in Paragraphs ---")
    found = False
    for i, p in enumerate(doc.paragraphs):
        if "{NOMBRE BENEFICIARIO}" in p.text or "NOMBRE BENEFICIARIO" in p.text:
            print(f"Match found in Paragraph {i}:")
            print(f"  Text: '{p.text}'")
            print("  Runs:")
            for r in p.runs:
                print(f"    [{r.text}]")
            found = True
            
    if not found:
        print("❌ Placeholder {NOMBRE BENEFICIARIO} NOT found in main body paragraphs.")


if __name__ == "__main__":
    inspect_template("CONTRATO 2026 sobrevivencia -  plantilla.docx")
