import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io
import google.generativeai as genai
from datetime import datetime

# --- IMPORTS PARA OCR ---
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
# -------------------------------

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(layout="wide", page_title="Asesor Previsional IA")

st.sidebar.info("🤖 Asistente de Asesoría Previsional IA")
st.sidebar.divider()
st.sidebar.subheader("Modificar Informe")

# Caja de texto para modificar el informe
instrucciones_mod = st.sidebar.text_area(
    "Indicaciones de Modificación",
    help="Usa esta caja para pedirle a la IA que refine el informe (ej. 'Elimina los puntos 1 y 2 de la nota', 'Acorta la sección 6', 'Cambia el tono a más formal').",
    key="instrucciones_mod"
)

# El botón "Refrescar" se definirá en la lógica principal
# -------------------------------


# --- 2. FUNCIONES DE LECTURA Y IA ---

@st.cache_data
def leer_pdfs_cargados(files):
    """
    Lee el texto de múltiples archivos PDF.
    Si una página parece escaneada, aplica OCR automáticamente.
    """
    contexto_completo = ""
    st.write("Archivos recibidos para análisis:")
    
    for file in files:
        st.caption(f"- {file.name}")
        try:
            full_text = ""
            # Abrir el PDF en memoria con PyMuPDF (fitz)
            doc = fitz.open(stream=io.BytesIO(file.read()), filetype="pdf")
            
            for i, page in enumerate(doc):
                page_num = i + 1
                
                # 1. Intentar extracción de texto digital
                text = page.get_text("text")
                
                # 2. Heurística: Si el texto es muy corto, probablemente es escaneado
                if len(text.strip()) < 150: # Umbral de 150 caracteres
                    st.warning(f"Página {page_num} de {file.name} parece escaneada. Iniciando OCR... (esto puede tardar)")
                    
                    # 3. Renderizar la página como imagen (300 DPI)
                    zoom = 300 / 72  # 300 DPI / 72 DPI (default)
                    mat = fitz.Matrix(zoom, zoom)
                    pix = page.get_pixmap(matrix=mat)
                    
                    # 4. Convertir a formato PIL (Pillow)
                    img_data = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_data))
                    
                    # 5. Usar Tesseract para OCR en español
                    try:
                        # 'spa' = Spanish
                        ocr_text = pytesseract.image_to_string(img, lang='spa')
                        full_text += f"\n\n--- PÁGINA {page_num} ({file.name}) [Texto extraído por OCR] ---\n\n{ocr_text}"
                    except Exception as ocr_error:
                        st.error(f"Error de OCR en página {page_num}. Asegúrate de que Tesseract esté instalado y 'spa' (español) esté disponible. Error: {ocr_error}")
                        full_text += f"\n\n--- PÁGINA {page_num} ({file.name}) [ERROR DE OCR] ---\n\n"
                
                else:
                    # Es un PDF digital, usar el texto extraído
                    full_text += f"\n\n--- PÁGINA {page_num} ({file.name}) [Texto digital] ---\n\n{text}"
            
            contexto_completo += f"\n\n=== INICIO DOCUMENTO: {file.name} ===\n{full_text}\n=== FIN DOCUMENTO: {file.name} ===\n\n"
            doc.close()
        
        except Exception as e:
            st.error(f"Error al leer {file.name}: {e}")
    return contexto_completo

# === PROMPT PASO 1: ANÁLISIS (SECCIONES 1-5) ===
PROMPT_ANALISIS = """
Eres un Asesor Previsional experto y senior, con profundo conocimiento del sistema de pensiones chileno (AFP, SCOMP, PGU, APV, etc.).
Tu tarea es analizar TODOS los documentos de antecedentes que te entregaré (SCOMP, Certificado de Saldo, etc.) y generar un **Informe de Análisis** que contenga ÚNICAMENTE las secciones 1 a 5.
REGLAS IMPORTANTES:
1.  **Actúa como un experto:** Tu tono debe ser profesional y claro.
2.  **Cíñete a los datos:** No inventes información. Si un dato no se encuentra en los documentos (ej. Fecha de Nacimiento), debes indicarlo explícitamente (ej: "Fecha de Nacimiento: No informada en los documentos").
3.  **Calcula cuando se pida:** Para las Rentas Vitalicias Aumentadas, DEBES calcular los montos aumentados (Pensión Aumentada UF/$, Pensión Líquida Aumentada) basándote en la "pensión base" que encuentres en el SCOMP.
4.  **Usa Markdown:** Estructura tu respuesta usando Markdown (títulos, negritas, tablas).
5.  **Fecha del Informe:** {FECHA_HOY}
6.  **NO INCLUYAS la Sección 6 (Recomendación Final).** Termina el informe después de la Sección 5.
7.  **Formato de Títulos:** Usa '##' para Secciones (ej. ## 1) Antecedentes) y '###' para Subsecciones (ej. ### Certificado de Saldos). Usa '####' para los títulos de las modalidades (ej. #### a) Retiro programado).
---
TEXTO EXTRAÍDO DE LOS DOCUMENTOS DEL CLIENTE (SCOMP, CARTOLAS, ETC.):
{CONTEXTO_DOCUMENTOS}
---
Basado ÚNICAMENTE en los documentos, genera el informe con la siguiente estructura exacta (Secciones 1 a 5):
## Informe final de Asesoría Previsional
### 1) Antecedentes del afiliado y certificado SCOMP
* **Nombre Completo:** [Extraer]
* **RUT:** [Extraer]
* **Fecha de Nacimiento:** [Extraer]
* **Edad Cumplida (a la fecha actual):** [Calcular o extraer si está]
* **Sexo:** [Extraer]
* **Estado Civil:** [Extraer]
* **AFP de Origen:** [Extraer]
* **Institución de Salud:** [Extraer o poner "No informada"]
* **Fecha Solicitud de Pensión:** [Extraer]
* **Fecha de Emisión Certificado de Ofertas (SCOMP):** [Extraer]
* **Período de Aceptación de Ofertas:** [Extraer fechas inicio y fin]
#### Certificado de Saldos
**Descripción:** El saldo total destinado a pensión (Cotizaciones Obligatorias, Fondo [Extraer Fondo]) es de **UF [Extraer Saldo UF]**. Este monto equivale a **$[Extraer Saldo $]**. El valor de la UF utilizado es de **$[Extraer Valor UF]** al **[Extraer Fecha UF]**. Este Certificado se encuentra vigente hasta el día **[Extraer Vigencia Saldo]**.
### 2) Antecedentes del beneficiario
[Extraer los datos del beneficiario en formato tabla o lista: Nombre, RUT, Parentesco. Si no existen, escribir: "El afiliado declara no contar con beneficiarios legales de pensión."]
### 3) Situación previsional
* **Tipo de Pensión Solicitada:** [Extraer, ej: Vejez Edad, Cambio de Modalidad]
* **Saldo para Pensión:** **UF [Extraer Saldo UF]**
* **Modalidades Solicitadas al SCOMP:** [Extraer las modalities que se pidieron, ej: RVIS, RVA 100% 36m]
### 4) Gestiones realizadas
[Describir las gestiones en formato lista o tabla, extrayendo fechas y acciones. Ej:
* **Solicitud de Pensión de Vejez Edad:** Presentada el [Fecha] a AFP [Nombre].
* **Retiro Certificado de Saldos:** Se retira el día [Fecha].
* **Solicitud de Ofertas (SCOMP):** Ingresada el [Fecha], por el Asesor Previsional [Nombre Asesor].]
### 5) Resultados Scomp
#### a) Retiro programado
**Descripción:** Es una modalidad de pensión que se paga con cargo a la Cuenta de Capitalización Individual del afiliado. La pensión se recalcula anualmente, considerando el saldo remanente, la expectativa de vida del afiliado y de sus beneficiarios, y la rentabilidad del fondo. Por lo tanto, la pensión puede subir o bajar cada año.
**Cuadro de resultados:**
[Generar tabla Markdown con TODAS las AFP del SCOMP]
| AFP | Pensión en UF | Pensión Bruta en $| Descuento 7% Salud$ | Descuento Comisión AFP $ | Pensión Líquida en $ |
| :--- | :--- | :--- | :--- | :--- | :--- |
| [AFP 1] | [uf] | [bruta] | [salud] | [comision] | [liquida] |
| [AFP 2] | [uf] | [bruta] | [salud] | [comision] | [liquida] |
| ... | ... | ... ... | ... | ... |
**Nota:** La oferta de Retiro Programado de su AFP de Origen ([Nombre AFP Origen]) es de **[UF] UF** al mes, lo que equivale a una Pensión Bruta de **$[Monto $]**. Con el descuento de salud ($[Monto Salud]) y la comisión de la AFP ($[Monto Comisión]), la pensión líquida aproximada es de **$[Monto Líquido]** para el primer año.
#### b) Renta Vitalicia
**Renta Vitalicia Inmediata Simple**
**Descripción:** Es un contrato con una Compañía de Seguros, donde el afiliado traspasa la totalidad de su saldo para recibir una pensión mensual en UF fija y de por vida. El monto no varía, independiente de la rentabilidad del mercado o de la expectativa de vida.
**Cuadro de resultados (4 mejores ofertas):**
| Compañía de Seguros | Pensión en UF | Pensión Bruta $| Descuento 7% Salud$ | Pensión Líquida $ |
| :--- | :--- | :--- | :--- | :--- |
| [Cia 1] | [uf] | [bruta] | [salud] | [liquida] |
| [Cia 2] | [uf] | [bruta] | [salud] | [liquida] |
| [Cia 3] | [uf] | [bruta] | [salud] | [liquida] |
| [Cia 4] | [uf] | [bruta] | [salud] | [liquida] |
**Renta Vitalicia Aumentada**
**Descripción:** La "Cláusula de Aumento Temporal de Pensión" es una cobertura adicional que permite duplicar (aumentar en un 100%) el monto de la pensión durante un período determinado al inicio. Una vez que este período finaliza, la pensión vuelve a su monto base original, el cual es fijo en UF y se paga de por vida.
[Generar una sección para CADA modalidad de Renta Vitalicia Aumentada encontrada en el SCOMP, ej: "Renta Vitalicia Aumentada 100% por 36 Meses"]
**[Título de la Modalidad, ej: Renta Vitalicia Aumentada 100% por 36 Meses]**
**Cuadro de resultados (4 mejores ofertas):**
| Compañía | Pensión Aumentada en UF | Pensión Aumentada en $| Descuento 7% Salud$ | Pensión Líquida Período Aumentado | Pensión Después de Aumento en UF (Base) |
| :--- | :--- | :--- | :--- | :--- | :--- |
| [Cia 1] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
| [Cia 2] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
| [Cia 3] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
| [Cia 4] | [Calcular: Base * 2] | [Calcular: Base $* 2] | [Calcular: (Base$ * 2) * 0.07] | [Calcular: (Base $ * 2) - Salud] | [Extraer Base UF] |
**Explicación:** Después del período aumentado, su pensión bajará al monto de la pensión base calculada. En este caso, la mejor oferta es de **[Base UF de la mejor oferta] UF**, lo que equivale a **$[Monto Base $]** brutos.
"""

# === PROMPT PASO 2: RECOMENDACIÓN (SECCIÓN 6) ===
PROMPT_RECOMENDACION = """
Eres un Asesor Previsional experto. Tu tarea es redactar la **Sección 6: Recomendación Final** para un informe.
Te entregaré el análisis de datos (Secciones 1-5) como contexto, y las instrucciones del asesor humano.
Redacta ÚNICAMENTE la "## 6) Recomendación Final" siguiendo las instrucciones.
---
INSTRUCCIONES DEL ASESOR HUMANO PARA LA RECOMENDACIÓN:
"{INSTRUCCIONES_USUARIO}"
---
CONTEXTO (ANÁLISIS DE DATOS SECCIONES 1-5):
{ANALISIS_PREVIO}
---
Redacta ÚNICAMENTE la "## 6) Recomendación Final":
"""

# === PROMPT PASO 3: MODIFICACIÓN ===
PROMPT_MODIFICACION = """
Eres un editor profesional. Tu tarea es tomar el siguiente informe previsional y modificarlo según las instrucciones del usuario.
REGLAS:
1.  **Aplica las modificaciones solicitadas** de forma precisa.
2.  **No cambies el formato Markdown** (títulos ##, ###, tablas |, etc.) a menos que la instrucción te lo pida.
3.  **Mantén el tono profesional** del informe.
4.  Entrega el **informe completo modificado**, no solo la parte que cambiaste.
---
INFORME ORIGINAL:
{INFORME_ACTUAL}
---
INSTRUCCIONES DEL USUARIO PARA MODIFICAR:
"{INSTRUCCIONES_MODIFICACION}"
---
INFORME MODIFICADO:
"""

@st.cache_data(show_spinner=False)
def generar_modificacion_ia(informe_actual, instrucciones, api_key):
    """
    Llama a la API de Gemini para MODIFICAR un informe ya existente.
    """
    if not api_key:
        st.error("API Key no configurada.")
        return None
    if not informe_actual or not instrucciones:
        st.error("Faltan datos para modificar el informe.")
        return None
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        prompt_completo = PROMPT_MODIFICACION.format(
            INFORME_ACTUAL=informe_actual,
            INSTRUCCIONES_MODIFICACION=instrucciones
        )
        
        generation_config = {"temperature": 0.2, "response_mime_type": "text/plain"}
        request_options = {"timeout": 300}
        
        response = model.generate_content(
            prompt_completo,
            generation_config=generation_config,
            request_options=request_options
        )
        return response.text
    except Exception as e:
        st.error(f"Error al modificar el informe con IA: {e}")
        st.exception(e)
        return None


@st.cache_data(show_spinner=False)
def generar_analisis_ia(contexto, api_key):
    """
    Llama a la API de Gemini para generar el ANÁLISIS (Secciones 1-5).
    """
    if not api_key:
        st.error("API Key no configurada.")
        return None
    if not contexto:
        st.error("Contexto de PDF vacío.")
        return None
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        fecha_hoy_str = datetime.now().strftime('%d de %B de %Y')
        prompt_completo = PROMPT_ANALISIS.format(
            CONTEXTO_DOCUMENTOS=contexto,
            FECHA_HOY=fecha_hoy_str
        )
        
        generation_config = {"temperature": 0.1, "response_mime_type": "text/plain"}
        request_options = {"timeout": 300} 
        
        response = model.generate_content(
            prompt_completo,
            generation_config=generation_config,
            request_options=request_options
        )
        return response.text
    except Exception as e:
        st.error(f"Error al generar el análisis con IA: {e}")
        st.exception(e)
        return None

@st.cache_data(show_spinner=False)
def generar_recomendacion_ia(analisis_previo, instrucciones, api_key):
    """
    Llama a la API de Gemini para generar SOLO la RECOMENDACIÓN (Sección 6).
    """
    if not api_key:
        st.error("API Key no configurada.")
        return None
    if not analisis_previo or not instrucciones:
        st.error("Faltan datos para generar la recomendación.")
        return None
        
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro-latest')
        
        prompt_completo = PROMPT_RECOMENDACION.format(
            ANALISIS_PREVIO=analisis_previo,
            INSTRUCCIONES_USUARIO=instrucciones
        )
        
        generation_config = {"temperature": 0.2, "response_mime_type": "text/plain"}
        request_options = {"timeout": 120}
        
        response = model.generate_content(
            prompt_completo,
            generation_config=generation_config,
            request_options=request_options
        )
        return response.text
    except Exception as e:
        st.error(f"Error al generar la recomendación con IA: {e}")
        st.exception(e)
        return None


# --- 3. FUNCIONES DE DESCARGA (SOLO DOCX) ---

def crear_reporte_doc(informe_texto):
    """
    Crea un archivo .docx en memoria, interpretando Markdown,
    con fuente "Roboto" y sin asteriscos.
    """
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Roboto'
    font.size = Pt(11)

    styles = doc.styles
    for h_level in [1, 2, 3, 4]:
        try:
            h_style = styles[f'Heading {h_level}']
            h_style.font.name = 'Roboto'
            h_style.font.bold = True
        except KeyError:
            pass
            
    try:
        bullet_style = styles['List Bullet']
        bullet_style.font.name = 'Roboto'
        bullet_style.font.size = Pt(11)
    except KeyError:
        pass

    in_table = False
    table = None
    
    for line in informe_texto.split('\n'):
        line_stripped = line.strip().replace('*', '')

        if line.strip().startswith('|') and line.strip().endswith('|'):
            cells = [c.strip().replace('*', '') for c in line.strip().split('|')[1:-1]]
            
            if '---' in cells[0]:
                continue

            if not in_table:
                try:
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, item in enumerate(cells):
                        hdr_cells[i].text = item
                        run = hdr_cells[i].paragraphs[0].runs[0]
                        run.font.name = 'Roboto'
                        run.font.bold = True
                    in_table = True
                except Exception as e:
                    st.warning(f"Error al crear cabecera de tabla DOCX: {e}")
            else:
                try:
                    row_cells = table.add_row().cells
                    for i, item in enumerate(cells):
                         if i < len(row_cells):
                            row_cells[i].text = item
                            run = row_cells[i].paragraphs[0].runs[0]
                            run.font.name = 'Roboto'
                except Exception as e:
                     st.warning(f"Error al añadir fila a tabla DOCX: {e}")
        
        else:
            if in_table:
                doc.add_paragraph() 
                in_table = False
                table = None

            if line.strip().startswith('## '):
                doc.add_heading(line_stripped.replace('## ', ''), level=2)
            elif line.strip().startswith('### '):
                doc.add_heading(line_stripped.replace('### ', ''), level=3)
            elif line.strip().startswith('#### '):
                doc.add_heading(line_stripped.replace('#### ', ''), level=4)
            elif line.strip().startswith('* '):
                doc.add_paragraph(line_stripped, style='List Bullet')
            elif line_stripped and not line_stripped.startswith('---'):
                p = doc.add_paragraph()
                p.add_run(line_stripped)

    if in_table:
        doc.add_paragraph()

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# --- 4. LÓGICA PRINCIPAL DE LA APLICACIÓN ---

st.title("🤖 Asistente de Asesoría Previsional (IA)")
st.write("Carga todos los documentos de tu cliente (SCOMP, Cartolas, APV, etc.) para generar un informe de asesoría consolidado.")

# --- Estados de Sesión ---
if 'contexto_documentos' not in st.session_state:
    st.session_state.contexto_documentos = None
# Esta es la ÚNICA variable que guarda el texto del informe
if 'informe_actual' not in st.session_state:
    st.session_state.informe_actual = None
# -------------------------

uploaded_files = st.file_uploader(
    "1. Cargar antecedentes del cliente (PDF)", 
    type=["pdf"],
    accept_multiple_files=True
)

st.divider()

# --- PASO 1: Generar Análisis (Secciones 1-5) ---
if uploaded_files:
    # Esta línea ahora solo se ejecuta si los archivos cambian
    with st.spinner("Leyendo y procesando los archivos PDF..."):
        st.session_state.contexto_documentos = leer_pdfs_cargados(uploaded_files)
    
    if st.button("Generar Análisis de Datos (Secciones 1-5)", type="primary"):
        try:
            final_api_key = st.secrets["api_key"]
        except:
            st.error("Error: La API Key no está configurada en los 'secrets' de la aplicación.")
            final_api_key = None
        
        if final_api_key and st.session_state.contexto_documentos:
            with st.spinner("La IA está analizando los datos (Secciones 1-5)... (Esto puede tardar hasta 1 minuto)"):
                analisis_resultado = generar_analisis_ia(
                    st.session_state.contexto_documentos,
                    final_api_key
                )
            
            if analisis_resultado:
                st.session_state.informe_actual = analisis_resultado # Guarda el análisis (1-5)
                st.success("Análisis (Secciones 1-5) generado. Ya puedes modificarlo o añadir la recomendación.")
            else:
                st.error("No se pudo generar el análisis.")
        elif not st.session_state.contexto_documentos:
             st.error("Error: No se pudo leer el contexto de los PDF.")

# --- Lógica de Refresco (Sidebar) ---
# Este botón ahora funciona si el informe_actual (Sec 1-5) existe
if st.sidebar.button("Refrescar Informe con Modificaciones"):
    if st.session_state.informe_actual and st.session_state.instrucciones_mod:
        try:
            final_api_key = st.secrets["api_key"]
        except:
            st.error("Error: API Key no configurada.")
            final_api_key = None
        
        if final_api_key:
            with st.spinner("La IA está aplicando tus modificaciones..."):
                informe_modificado = generar_modificacion_ia(
                    st.session_state.informe_actual, # Envía el informe actual (1-5 o 1-6)
                    st.session_state.instrucciones_mod,
                    final_api_key
                )
            if informe_modificado:
                st.session_state.informe_actual = informe_modificado # Sobrescribe el informe
                st.success("Informe refrescado.")
            else:
                st.error("No se pudo modificar el informe.")
    elif not st.session_state.informe_actual:
        st.sidebar.warning("Debes generar el 'Análisis de Datos' (Sección 1-5) primero.")
    else:
        st.sidebar.warning("Escribe alguna instrucción de modificación en la caja de texto.")


# --- PASO 2 y 3: Mostrar Informe, Pedir Recomendación y Descargar ---
# [INICIO BLOQUE CORREGIDO]
# Esta sección ahora se muestra solo si el informe_actual existe.
if st.session_state.informe_actual:
    
    st.subheader("Vista Previa del Informe Actual")
    st.markdown(st.session_state.informe_actual)
    
    st.divider()
    st.subheader("2. Instrucciones para la Recomendación Final (Sección 6)")
    # Este widget se dibuja aquí
    st.text_area(
        "Escriba sus instrucciones para la recomendación:", 
        key="instrucciones_rec", 
        height=150,
        help="Escribe aquí tus ideas (ej. 'Recomendar RVA a 60m...') y presiona el botón de abajo para AÑADIR la Sección 6 al informe."
    )

    if st.button("Añadir Recomendación al Informe (Sección 6)", type="primary"):
        
        try:
            final_api_key = st.secrets["api_key"]
        except:
            st.error("Error: La API Key no está configurada en los 'secrets' de la aplicación.")
            final_api_key = None

        instrucciones_texto = st.session_state.instrucciones_rec
        
        if final_api_key and instrucciones_texto:
            with st.spinner("La IA está redactando y añadiendo la recomendación (Sección 6)..."):
                recomendacion_resultado = generar_recomendacion_ia(
                    st.session_state.informe_actual, # Usa el informe actual (1-5) como contexto
                    instrucciones_texto,
                    final_api_key
                )
            
            if recomendacion_resultado:
                # --- Lógica de AÑADIR ---
                st.session_state.informe_actual += "\n\n" + recomendacion_resultado
                
                # --- Limpiar la caja de texto para evitar duplicados ---
                # ESTA LÍNEA ESTÁ AHORA DENTRO DEL IF, ESTA ES LA CORRECCIÓN
                st.session_state.instrucciones_rec = "" 
                
                st.success("Recomendación añadida. Ya puedes modificar el informe completo o descargarlo.")
                
                # --- LA CORRECCIÓN ---
                st.rerun() # <-- Esto aplica la limpieza de la caja de texto
                
            else:
                st.error("No se pudo generar la recomendación.")
        elif not instrucciones_texto:
            st.warning("Por favor, escriba las instrucciones para la recomendación.")

    # --- Sección de Descarga ---
    st.divider()
    st.subheader("Descargar Informe Completo")
    
    try:
        informe_completo_texto = st.session_state.informe_actual
        
        doc_data = crear_reporte_doc(informe_completo_texto)
        
        st.download_button(
            label="📄 Descargar Informe en DOCX (Compatible con Word/Google Docs)",
            data=doc_data,
            file_name="Informe_final_Asesoria_Previsional.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
            
    except Exception as e:
        st.error(f"Error al generar el archivo de descarga: {e}")
        st.exception(e)
# [FIN BLOQUE CORREGIDO]

# La línea de error "st.session_state.instrucciones_rec = """ NO DEBE ESTAR AQUÍ.
# Como puedes ver, no está en este código.
